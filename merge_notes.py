from docx import Document
from docx.shared import Inches
import os
import re

def convert_to_seconds(timestamp):
    time_parts = re.split('[:.]', timestamp)
    seconds = int(time_parts[0]) * 3600 + int(time_parts[1]) * 60 + float(time_parts[2])
    return seconds

def add_frames_to_docx(frames_folder, transcript_file, output_docx='new_output.docx'):
    document = Document()
    
    # Load transcript and timestamps
    with open(transcript_file, 'r') as f:
        transcript_lines = f.readlines()

    # Sort frames by timestamp (assumed to be in milliseconds)
    frames = sorted(os.listdir(frames_folder), key=lambda x: int(os.path.splitext(x)[0]))
    frame_times = [int(os.path.splitext(frame)[0]) / 1000.0 for frame in frames]  # Convert milliseconds to seconds

    frame_idx = 0
    total_frames = len(frames)

    for i, line in enumerate(transcript_lines):
        match = re.match(r'(\d{2}:\d{2}:\d{2}\.\d{3}) --> (\d{2}:\d{2}:\d{2}\.\d{3})', line)
        if match:
            start_time = convert_to_seconds(match.group(1))
            end_time = convert_to_seconds(match.group(2))

            # Insert all frames that fit within this transcript's time range
            while frame_idx < total_frames and frame_times[frame_idx] < start_time:
                frame_idx += 1

            # Insert corresponding frames
            while frame_idx < total_frames and start_time <= frame_times[frame_idx] <= end_time:
                frame_path = os.path.join(frames_folder, frames[frame_idx])
                document.add_picture(frame_path, width=Inches(2))
                frame_idx += 1

            # Insert the transcript text
            transcript = transcript_lines[i + 1].strip() if i + 1 < len(transcript_lines) else ""
            document.add_paragraph(transcript)
    
    # Save the final document
    document.save(output_docx)

# frames_folder = 'frames'
# transcript_file = 'timestamps.txt'
# add_frames_to_docx(frames_folder, transcript_file)

def main():
    frames_folder = 'frames'
    transcript_file = 'translated_timestamps.txt'
    add_frames_to_docx(frames_folder, transcript_file)